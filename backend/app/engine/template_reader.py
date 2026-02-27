"""Read Word template files and extract available styles."""
from __future__ import annotations

import io

from docx import Document


def extract_styles(file_bytes: bytes) -> list[str]:
    """Extract all paragraph and character style names from a .docx template."""
    doc = Document(io.BytesIO(file_bytes))
    styles = []
    for style in doc.styles:
        if style.type in (1, 2):  # 1=paragraph, 2=character
            if style.name and not style.name.startswith("_"):
                styles.append(style.name)
    return sorted(set(styles))


def validate_styles(file_bytes: bytes, required_styles: list[str]) -> dict:
    """Check which required styles exist in the template."""
    available = set(extract_styles(file_bytes))
    found = []
    missing = []
    for style in required_styles:
        if style in available:
            found.append(style)
        else:
            missing.append(style)
    return {"found": found, "missing": missing}


def _collect_paragraphs(doc) -> list:
    """Collect all paragraphs from body, tables, headers, and footers in document order."""
    paragraphs = []

    # Body paragraphs
    for para in doc.paragraphs:
        paragraphs.append(("body", para))

    # Table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    paragraphs.append(("table", para))

    # Header and footer paragraphs
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header and header.is_linked_to_previous is False:
                for para in header.paragraphs:
                    paragraphs.append(("header", para))
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer and footer.is_linked_to_previous is False:
                for para in footer.paragraphs:
                    paragraphs.append(("footer", para))

    return paragraphs


def extract_used_styles(file_bytes: bytes) -> list[str]:
    """Extract style names actually applied to content in the document."""
    doc = Document(io.BytesIO(file_bytes))
    used = set()

    for _source, para in _collect_paragraphs(doc):
        if para.style and para.style.name:
            used.add(para.style.name)
        for run in para.runs:
            if run.style and run.style.name and run.style.name != "Default Paragraph Font":
                used.add(run.style.name)

    return sorted(used)


def extract_style_map(file_bytes: bytes) -> list[dict]:
    """Build an ordered list of paragraph style assignments for the style map overlay."""
    doc = Document(io.BytesIO(file_bytes))
    result = []
    idx = 0

    for source, para in _collect_paragraphs(doc):
        text = para.text.strip()
        style_name = para.style.name if para.style and para.style.name else "Normal"
        preview = text[:120] if text else ""
        result.append({
            "index": idx,
            "text": preview,
            "style": style_name,
        })
        idx += 1

    return result
