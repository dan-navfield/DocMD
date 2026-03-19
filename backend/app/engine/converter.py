"""Core conversion engine: Markdown + Template + Mapping → Word document."""
from __future__ import annotations

import io
import logging
from typing import Any

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

from app.engine.parser import parse_markdown
from app.engine.template_reader import extract_styles
from app.engine.validator import validate_mapping, check_unmapped_elements

logger = logging.getLogger(__name__)

# Word-processing XML namespace
WP_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


class MarkdownToWordConverter:
    """Converts Markdown to a styled Word document using mapping rules."""

    def convert(
        self,
        markdown_text: str,
        template_bytes: bytes,
        mapping_rules: dict,
    ) -> tuple[bytes, dict]:
        """
        Convert Markdown text to a Word document.

        Returns:
            tuple of (docx_bytes, conversion_report)
        """
        # Parse markdown to AST
        ast = parse_markdown(markdown_text)

        # Load template
        doc = Document(io.BytesIO(template_bytes))
        available_styles = extract_styles(template_bytes)

        # Validate mapping against template
        style_warnings = validate_mapping(mapping_rules, available_styles)
        unmapped_warnings = check_unmapped_elements(ast, mapping_rules)

        # Prepare template: preserve cover page + final page, clear body sections
        saved_tail = self._prepare_template(doc)

        # Try to update the cover page title with the first H1 from the markdown
        first_h1_node = None
        for node in ast:
            if node.get("type") == "heading" and node.get("level") == 1:
                first_h1_node = node
                break

        cover_title_updated = False
        if first_h1_node and mapping_rules.get("document_title"):
            h1_text = self._extract_plain_text(first_h1_node.get("children", []))
            cover_title_updated = self._update_cover_title(
                doc, h1_text, mapping_rules
            )

        # Clear/replace cover page metadata fields (subtitle, organization, etc.)
        self._update_cover_metadata(doc, mapping_rules)

        # Stats tracking
        stats = {
            "headings": 0,
            "paragraphs": 0,
            "lists": 0,
            "tables": 0,
            "code_blocks": 0,
            "images": 0,
            "_first_h1_done": cover_title_updated,
        }

        # Walk AST and build document
        for node in ast:
            # Skip first H1 if it was used to update the cover page title
            if cover_title_updated and node is first_h1_node:
                stats["headings"] += 1
                continue
            self._process_node(doc, node, mapping_rules, stats, level=0)

        # Re-insert saved tail (last section break + final page content)
        if saved_tail:
            body_sectpr = doc.element.body.find(qn("w:sectPr"))
            for elem in saved_tail:
                if body_sectpr is not None:
                    body_sectpr.addprevious(elem)
                else:
                    doc.element.body.append(elem)

        # Generate output bytes
        output = io.BytesIO()
        doc.save(output)
        docx_bytes = output.getvalue()

        # Build report (exclude internal stats prefixed with _)
        all_warnings = style_warnings + unmapped_warnings
        public_stats = {k: v for k, v in stats.items() if not k.startswith("_")}
        report = {
            "elements_processed": sum(public_stats.values()),
            "warnings": all_warnings,
            "warnings_text": [
                f"{w['type']}: {w.get('style', w.get('element', ''))}"
                for w in all_warnings
            ],
            "stats": public_stats,
        }

        return docx_bytes, report

    # ------------------------------------------------------------------
    # Template preparation
    # ------------------------------------------------------------------

    def _prepare_template(self, doc: Document) -> list:
        """Prepare template body for content insertion.

        Identifies section boundaries in the template and preserves:
        - First section (cover page) with its section break
        - Last section (final page) content and body-level sectPr
        - The section break before the final page (for header/footer continuity)

        Middle section content and extra breaks are removed so converted
        markdown can be inserted into a clean body section.

        Returns a list of saved XML elements (last section break paragraph +
        final page content) that should be re-inserted after converted
        content is added.  The caller inserts them before the body-level
        sectPr to reconstruct the three-section layout:
            [cover] [sectPr1] [content…] [sectPrN-1] [final page] [body sectPr]
        """
        body = doc.element.body
        all_elements = list(body)

        # Find section-break paragraphs (w:p elements whose w:pPr contains w:sectPr)
        section_breaks = []
        for i, elem in enumerate(all_elements):
            if elem.tag == f"{WP_NS}p":
                pPr = elem.find(f"{WP_NS}pPr")
                if pPr is not None and pPr.find(f"{WP_NS}sectPr") is not None:
                    section_breaks.append((elem, i))

        num_breaks = len(section_breaks)
        logger.info(
            "Template has %d section break(s) → %d section(s)",
            num_breaks,
            num_breaks + 1,
        )

        # --- Single section (no breaks): clear everything, keep body sectPr ---
        if num_breaks == 0:
            for element in list(body):
                if element.tag.endswith("}sectPr"):
                    continue
                body.remove(element)
            return []

        # --- Two sections (cover + body): keep cover, clear body ---
        if num_breaks == 1:
            first_break_idx = section_breaks[0][1]
            to_remove = []
            for i, elem in enumerate(all_elements):
                if i <= first_break_idx:
                    continue  # keep cover page + its break
                if elem.tag.endswith("}sectPr"):
                    continue  # keep body-level sectPr
                to_remove.append(elem)
            for elem in to_remove:
                body.remove(elem)
            logger.info("2-section template: preserved cover, cleared body section")
            return []

        # --- 3+ sections: keep cover, save last break + final page, clear middle ---
        first_break_idx = section_breaks[0][1]
        last_break_idx = section_breaks[-1][1]

        saved_tail = []
        to_remove = []

        for i, elem in enumerate(all_elements):
            # Keep cover section intact (everything up to and including first break)
            if i <= first_break_idx:
                continue
            # Always keep the body-level sectPr (direct child of body)
            if elem.tag.endswith("}sectPr"):
                continue
            if i >= last_break_idx:
                # Save for re-insertion: last section break + final page content
                saved_tail.append(elem)
                to_remove.append(elem)
            else:
                # Middle section content and extra section breaks: remove
                to_remove.append(elem)

        for elem in to_remove:
            body.remove(elem)

        logger.info(
            "Multi-section template: preserved cover (section 1), "
            "cleared %d middle section(s), saved %d final-page element(s)",
            num_breaks - 1,
            len(saved_tail),
        )
        return saved_tail

    def _update_cover_title(
        self, doc: Document, title_text: str, mapping_rules: dict
    ) -> bool:
        """Find the document-title paragraph on the cover page and update its text.

        Looks for a paragraph whose style matches the ``document_title``
        mapping rule (e.g. "Cover heading").  If found, replaces the text
        while preserving the original run formatting.

        Returns True if a matching paragraph was found and updated.
        """
        doc_title_style = mapping_rules.get("document_title", "")
        if not doc_title_style:
            return False

        for para in doc.paragraphs:
            if para.style and para.style.name == doc_title_style:
                # Replace text in the first run, clear the rest
                if para.runs:
                    para.runs[0].text = title_text
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    para.add_run(title_text)
                logger.info(
                    "Updated cover title (style=%r) → %r",
                    doc_title_style,
                    title_text[:80],
                )
                return True

        logger.info(
            "No paragraph with style %r found on cover page — "
            "first H1 will be added to body with that style instead",
            doc_title_style,
        )
        return False

    def _update_cover_metadata(
        self, doc: Document, mapping_rules: dict
    ) -> None:
        """Clear non-title text on the cover page.

        Two modes:
        1. **Automatic** (default): clears ALL text paragraphs in the cover
           section (everything before the first section break) except the
           paragraph matched by ``document_title``.  This handles any
           template without needing explicit per-field mapping.
        2. **Explicit**: if ``metadata_mapping`` is populated, only the
           paragraphs whose styles match the mapped values are cleared.
           ``document_subtitle`` is treated as a shorthand for
           ``metadata_mapping["subtitle"]``.

        Images and backgrounds embedded in headers/footers are unaffected.
        """
        doc_title_style = mapping_rules.get("document_title", "")
        meta = dict(mapping_rules.get("metadata_mapping") or {})

        # Honour the document_subtitle convenience field
        doc_subtitle_style = mapping_rules.get("document_subtitle", "")
        if doc_subtitle_style and "subtitle" not in meta:
            meta["subtitle"] = doc_subtitle_style

        # --- Explicit mode: clear only specifically-mapped styles ---
        if meta:
            for field_name, style_name in meta.items():
                if not style_name:
                    continue
                for para in doc.paragraphs:
                    if para.style and para.style.name == style_name:
                        for run in para.runs:
                            run.text = ""
                        logger.info(
                            "Cleared cover metadata field %r (style=%r)",
                            field_name,
                            style_name,
                        )
                        break
            return

        # --- Automatic mode: clear all cover-section text except title ---
        body = doc.element.body
        first_break_idx = None

        for i, elem in enumerate(body):
            if elem.tag == f"{WP_NS}p":
                pPr = elem.find(f"{WP_NS}pPr")
                if pPr is not None and pPr.find(f"{WP_NS}sectPr") is not None:
                    first_break_idx = i
                    break

        if first_break_idx is None:
            return  # No section break → no distinct cover page

        cleared = 0
        for i, elem in enumerate(body):
            if i >= first_break_idx:
                break
            if elem.tag != f"{WP_NS}p":
                continue

            # Build a python-docx Paragraph wrapper for this element
            from docx.text.paragraph import Paragraph
            para = Paragraph(elem, doc.part)

            # Skip the title paragraph
            if doc_title_style and para.style and para.style.name == doc_title_style:
                continue

            # Skip paragraphs that have no visible text (spacers, etc.)
            if not para.text.strip():
                continue

            # Clear all runs
            for run in para.runs:
                run.text = ""
            cleared += 1

        if cleared:
            logger.info(
                "Auto-cleared %d cover page paragraph(s) (kept title style=%r)",
                cleared,
                doc_title_style,
            )

    def _extract_plain_text(self, children: list[dict]) -> str:
        """Extract plain text from inline content nodes."""
        parts = []
        for child in children:
            if child.get("type") == "text":
                parts.append(child.get("text", ""))
            elif "children" in child:
                parts.append(self._extract_plain_text(child["children"]))
        return "".join(parts)

    # ------------------------------------------------------------------
    # AST node processors
    # ------------------------------------------------------------------

    def _process_node(
        self,
        doc: Document,
        node: dict,
        rules: dict,
        stats: dict,
        level: int = 0,
    ):
        """Process a single AST node and add it to the document."""
        node_type = node.get("type", "")

        if node_type == "heading":
            self._add_heading(doc, node, rules, stats)
        elif node_type == "paragraph":
            self._add_paragraph(doc, node, rules, stats)
        elif node_type == "list":
            self._add_list(doc, node, rules, stats, level)
        elif node_type == "code_block":
            self._add_code_block(doc, node, rules, stats)
        elif node_type == "blockquote":
            self._add_blockquote(doc, node, rules, stats)
        elif node_type == "table":
            self._add_table(doc, node, rules, stats)
        elif node_type == "thematic_break":
            self._add_page_break(doc)

    def _add_heading(self, doc: Document, node: dict, rules: dict, stats: dict):
        heading_level = node.get("level", 1)
        heading_rules = rules.get("heading", {})

        # Check for page break before this heading level
        page_break_before = rules.get("page_break_before", [])
        if f"heading.{heading_level}" in page_break_before:
            self._add_page_break(doc)

        # If this is the first H1 and document_title style is set, use that instead
        document_title_style = rules.get("document_title", "")
        if heading_level == 1 and document_title_style and not stats.get("_first_h1_done"):
            style_name = document_title_style
            stats["_first_h1_done"] = True
        elif isinstance(heading_rules, dict):
            style_name = heading_rules.get(str(heading_level), f"Heading {heading_level}")
        else:
            style_name = f"Heading {heading_level}"

        para = doc.add_paragraph()
        self._try_set_style(doc, para, style_name)
        self._add_inline_content(para, node.get("children", []))
        stats["headings"] += 1

    def _add_paragraph(self, doc: Document, node: dict, rules: dict, stats: dict):
        style_name = rules.get("paragraph", "Normal")
        para = doc.add_paragraph()
        self._try_set_style(doc, para, style_name)
        self._add_inline_content(para, node.get("children", []))
        stats["paragraphs"] += 1

    def _add_list(self, doc: Document, node: dict, rules: dict, stats: dict, level: int = 0):
        ordered = node.get("ordered", False)

        for item in node.get("items", []):
            # Get style based on list type and nesting level
            if ordered:
                style_keys = ["list_ordered", "list_ordered_2", "list_ordered_3"]
            else:
                style_keys = ["list_bullet", "list_bullet_2", "list_bullet_3"]

            style_key = style_keys[min(level, len(style_keys) - 1)]
            style_name = rules.get(style_key, style_key.replace("_", " ").title())

            # Add the list item text
            for child in item.get("children", []):
                if child.get("type") == "paragraph":
                    para = doc.add_paragraph()
                    self._try_set_style(doc, para, style_name)
                    self._add_inline_content(para, child.get("children", []))

            # Handle nested lists
            for nested_list in item.get("nested_lists", []):
                if nested_list:
                    self._add_list(doc, nested_list, rules, stats, level + 1)

        stats["lists"] += 1

    def _add_code_block(self, doc: Document, node: dict, rules: dict, stats: dict):
        style_name = rules.get("code_block", "Code")
        text = node.get("text", "")

        para = doc.add_paragraph()
        self._try_set_style(doc, para, style_name)
        run = para.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        stats["code_blocks"] += 1

    def _add_blockquote(self, doc: Document, node: dict, rules: dict, stats: dict):
        style_name = rules.get("blockquote", "Quote")
        for child in node.get("children", []):
            if child.get("type") == "paragraph":
                para = doc.add_paragraph()
                self._try_set_style(doc, para, style_name)
                self._add_inline_content(para, child.get("children", []))

    def _add_table(self, doc: Document, node: dict, rules: dict, stats: dict):
        rows = node.get("rows", [])
        if not rows:
            return

        table_rules = rules.get("table", {})
        num_cols = max(len(row.get("cells", [])) for row in rows)
        num_rows = len(rows)

        table = doc.add_table(rows=num_rows, cols=num_cols)

        # Try to apply table style
        table_style = table_rules.get("style", "Table Grid") if isinstance(table_rules, dict) else "Table Grid"
        try:
            table.style = table_style
        except (KeyError, ValueError):
            pass

        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row.get("cells", [])):
                if j < num_cols:
                    cell = table.rows[i].cells[j]
                    cell.text = cell_text
                    # Bold header rows
                    if row.get("is_header"):
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

        stats["tables"] += 1

    def _add_page_break(self, doc: Document):
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_break(WD_BREAK.PAGE)

    def _add_inline_content(self, para, children: list[dict]):
        """Add inline content (text, bold, italic, code, links) to a paragraph."""
        for child in children:
            child_type = child.get("type", "")

            if child_type == "text":
                para.add_run(child.get("text", ""))

            elif child_type == "strong":
                for sub in child.get("children", []):
                    run = para.add_run(sub.get("text", ""))
                    run.bold = True

            elif child_type == "emphasis":
                for sub in child.get("children", []):
                    run = para.add_run(sub.get("text", ""))
                    run.italic = True

            elif child_type == "code_span":
                run = para.add_run(child.get("text", ""))
                run.font.name = "Courier New"
                run.font.size = Pt(9)

            elif child_type == "link":
                url = child.get("url", "")
                link_text = "".join(
                    sub.get("text", "") for sub in child.get("children", [])
                )
                run = para.add_run(link_text)
                run.underline = True
                # Add hyperlink (simplified - full hyperlink requires XML manipulation)
                run.font.color.rgb = None  # Will use theme color

            elif child_type == "image":
                # Images would need to be downloaded and inserted
                alt = child.get("alt", "Image")
                para.add_run(f"[{alt}]")

            elif child_type == "linebreak":
                para.add_run().add_break()

    def _try_set_style(self, doc: Document, para, style_name: str):
        """Try to set a paragraph style, falling back to Normal if not found.

        We iterate style objects directly instead of using the string-based
        lookup because python-docx's BabelFish lowercases built-in names
        (e.g. "heading 1") which breaks with templates that use non-standard
        styleIds (common with InDesign-exported DOCX files).
        """
        from docx.enum.style import WD_STYLE_TYPE

        # Direct object lookup — bypasses BabelFish alias table
        for s in doc.styles:
            if s.name == style_name and s.type == WD_STYLE_TYPE.PARAGRAPH:
                para.style = s
                return

        # Fallback: try the standard string-based lookup
        try:
            para.style = style_name
        except (KeyError, ValueError):
            pass
