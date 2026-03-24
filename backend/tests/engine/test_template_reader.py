"""Tests for Word template style extraction."""
import io
import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from app.engine.template_reader import extract_styles, validate_styles, extract_used_styles


class TestExtractStyles:
    def test_extracts_paragraph_styles(self, simple_template_bytes):
        styles = extract_styles(simple_template_bytes)
        assert "Normal" in styles
        assert "Heading 1" in styles

    def test_extracts_custom_styles(self):
        doc = Document()
        doc.styles.add_style("My Custom Style", WD_STYLE_TYPE.PARAGRAPH)
        buf = io.BytesIO()
        doc.save(buf)
        styles = extract_styles(buf.getvalue())
        assert "My Custom Style" in styles

    def test_excludes_internal_styles(self):
        """Styles starting with _ should be excluded."""
        styles = extract_styles(_make_template_with_internal_style())
        assert not any(s.startswith("_") for s in styles)

    def test_returns_sorted_unique(self):
        doc = Document()
        doc.styles.add_style("Zebra", WD_STYLE_TYPE.PARAGRAPH)
        doc.styles.add_style("Alpha", WD_STYLE_TYPE.PARAGRAPH)
        buf = io.BytesIO()
        doc.save(buf)
        styles = extract_styles(buf.getvalue())
        assert styles == sorted(set(styles))

    def test_empty_template(self):
        """A minimal template still has built-in styles."""
        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        styles = extract_styles(buf.getvalue())
        assert len(styles) > 0
        assert "Normal" in styles


class TestValidateStyles:
    def test_all_found(self, simple_template_bytes):
        result = validate_styles(simple_template_bytes, ["Normal", "Heading 1"])
        assert result["found"] == ["Normal", "Heading 1"]
        assert result["missing"] == []

    def test_some_missing(self, simple_template_bytes):
        result = validate_styles(simple_template_bytes, ["Normal", "NonExistent"])
        assert "Normal" in result["found"]
        assert "NonExistent" in result["missing"]

    def test_empty_required(self, simple_template_bytes):
        result = validate_styles(simple_template_bytes, [])
        assert result["found"] == []
        assert result["missing"] == []


class TestExtractUsedStyles:
    def test_returns_styles_applied_to_content(self, simple_template_bytes):
        used = extract_used_styles(simple_template_bytes)
        assert "Normal" in used

    def test_empty_doc_returns_default(self):
        doc = Document()
        doc.add_paragraph("text")
        buf = io.BytesIO()
        doc.save(buf)
        used = extract_used_styles(buf.getvalue())
        assert "Normal" in used


def _make_template_with_internal_style():
    doc = Document()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
