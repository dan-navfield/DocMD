"""Tests for the Markdown → Word converter."""
import io
import pytest
from docx import Document

from app.engine.converter import MarkdownToWordConverter
from tests.conftest import read_docx_paragraphs, read_docx_tables


@pytest.fixture
def converter():
    return MarkdownToWordConverter()


class TestBasicConversion:
    def test_returns_bytes_and_report(self, converter, simple_template_bytes, basic_mapping_rules):
        docx_bytes, report = converter.convert("# Hello", simple_template_bytes, basic_mapping_rules)
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0
        assert isinstance(report, dict)
        assert "stats" in report
        assert "warnings" in report

    def test_output_is_valid_docx(self, converter, simple_template_bytes, basic_mapping_rules):
        docx_bytes, _ = converter.convert("# Hello\n\nWorld", simple_template_bytes, basic_mapping_rules)
        doc = Document(io.BytesIO(docx_bytes))
        assert len(doc.paragraphs) > 0

    def test_empty_markdown(self, converter, simple_template_bytes, basic_mapping_rules):
        docx_bytes, report = converter.convert("", simple_template_bytes, basic_mapping_rules)
        assert isinstance(docx_bytes, bytes)
        assert report["stats"]["headings"] == 0
        assert report["stats"]["paragraphs"] == 0


class TestHeadingConversion:
    def test_heading_uses_mapped_style(self, converter, simple_template_bytes, basic_mapping_rules):
        docx_bytes, _ = converter.convert("# Title", simple_template_bytes, basic_mapping_rules)
        paras = read_docx_paragraphs(docx_bytes)
        heading_paras = [p for p in paras if "Title" in p["text"]]
        assert len(heading_paras) > 0

    def test_heading_levels_mapped_correctly(self, converter, simple_template_bytes, basic_mapping_rules):
        md = "# H1\n\n## H2\n\n### H3"
        docx_bytes, report = converter.convert(md, simple_template_bytes, basic_mapping_rules)
        assert report["stats"]["headings"] == 3

    def test_heading_stats_counted(self, converter, simple_template_bytes, basic_mapping_rules):
        md = "# One\n\n## Two"
        _, report = converter.convert(md, simple_template_bytes, basic_mapping_rules)
        assert report["stats"]["headings"] == 2


class TestParagraphConversion:
    def test_paragraph_uses_mapped_style(self, converter, simple_template_bytes, basic_mapping_rules):
        docx_bytes, _ = converter.convert("Hello world", simple_template_bytes, basic_mapping_rules)
        paras = read_docx_paragraphs(docx_bytes)
        hello_paras = [p for p in paras if "Hello world" in p["text"]]
        assert len(hello_paras) > 0
        assert hello_paras[0]["style"] == "Normal"

    def test_paragraph_stats_counted(self, converter, simple_template_bytes, basic_mapping_rules):
        _, report = converter.convert("First\n\nSecond", simple_template_bytes, basic_mapping_rules)
        assert report["stats"]["paragraphs"] == 2


class TestListConversion:
    def test_unordered_list(self, converter, simple_template_bytes, basic_mapping_rules):
        md = "- item one\n- item two"
        docx_bytes, report = converter.convert(md, simple_template_bytes, basic_mapping_rules)
        assert report["stats"]["lists"] >= 1
        paras = read_docx_paragraphs(docx_bytes)
        texts = [p["text"] for p in paras]
        assert "item one" in texts
        assert "item two" in texts

    def test_ordered_list(self, converter, simple_template_bytes, basic_mapping_rules):
        md = "1. first\n2. second"
        _, report = converter.convert(md, simple_template_bytes, basic_mapping_rules)
        assert report["stats"]["lists"] >= 1

    def test_nested_list(self, converter, simple_template_bytes, basic_mapping_rules):
        md = "- parent\n  - child"
        docx_bytes, _ = converter.convert(md, simple_template_bytes, basic_mapping_rules)
        paras = read_docx_paragraphs(docx_bytes)
        texts = [p["text"] for p in paras]
        assert "parent" in texts
        assert "child" in texts


class TestCodeBlockConversion:
    def test_code_block_rendered(self, converter, simple_template_bytes, basic_mapping_rules):
        md = "```\nprint('hello')\n```"
        docx_bytes, report = converter.convert(md, simple_template_bytes, basic_mapping_rules)
        assert report["stats"]["code_blocks"] == 1
        paras = read_docx_paragraphs(docx_bytes)
        code_texts = [p["text"] for p in paras if "print" in p["text"]]
        assert len(code_texts) > 0


class TestTableConversion:
    def test_table_rendered(self, converter, simple_template_bytes, basic_mapping_rules):
        md = "| A | B |\n|---|---|\n| 1 | 2 |"
        docx_bytes, report = converter.convert(md, simple_template_bytes, basic_mapping_rules)
        assert report["stats"]["tables"] == 1
        tables = read_docx_tables(docx_bytes)
        assert len(tables) >= 1
        assert tables[-1][0] == ["A", "B"]
        assert tables[-1][1] == ["1", "2"]


class TestTableCellFormatting:
    def test_bold_in_table_cell(self, converter, simple_template_bytes, basic_mapping_rules):
        md = "| **bold** | normal |\n|---|---|\n| a | b |"
        docx_bytes, _ = converter.convert(md, simple_template_bytes, basic_mapping_rules)
        doc = Document(io.BytesIO(docx_bytes))
        for table in doc.tables:
            cell = table.rows[0].cells[0]
            bold_runs = [r for r in cell.paragraphs[0].runs if r.bold]
            assert len(bold_runs) > 0, "Expected bold run in header cell"
            return
        pytest.fail("No table found")

    def test_link_in_table_cell(self, converter, simple_template_bytes, basic_mapping_rules):
        md = "| [link](https://x.com) |\n|---|\n| data |"
        docx_bytes, _ = converter.convert(md, simple_template_bytes, basic_mapping_rules)
        doc = Document(io.BytesIO(docx_bytes))
        from docx.oxml.ns import qn
        for table in doc.tables:
            cell_xml = table.rows[0].cells[0]._tc.xml
            assert "hyperlink" in cell_xml, "Expected hyperlink in table cell"
            return
        pytest.fail("No table found")


class TestBlockquoteConversion:
    def test_blockquote_rendered(self, converter, simple_template_bytes, basic_mapping_rules):
        md = "> quoted text"
        docx_bytes, _ = converter.convert(md, simple_template_bytes, basic_mapping_rules)
        paras = read_docx_paragraphs(docx_bytes)
        quoted = [p for p in paras if "quoted text" in p["text"]]
        assert len(quoted) > 0


class TestInlineFormatting:
    def test_bold_text_in_paragraph(self, converter, simple_template_bytes, basic_mapping_rules):
        md = "Hello **bold** world"
        docx_bytes, _ = converter.convert(md, simple_template_bytes, basic_mapping_rules)
        doc = Document(io.BytesIO(docx_bytes))
        for para in doc.paragraphs:
            if "bold" in para.text:
                bold_runs = [r for r in para.runs if r.bold and "bold" in r.text]
                assert len(bold_runs) > 0
                return
        pytest.fail("No paragraph containing 'bold' found")

    def test_italic_text_in_paragraph(self, converter, simple_template_bytes, basic_mapping_rules):
        md = "Hello *italic* world"
        docx_bytes, _ = converter.convert(md, simple_template_bytes, basic_mapping_rules)
        doc = Document(io.BytesIO(docx_bytes))
        for para in doc.paragraphs:
            if "italic" in para.text:
                italic_runs = [r for r in para.runs if r.italic and "italic" in r.text]
                assert len(italic_runs) > 0
                return
        pytest.fail("No paragraph containing 'italic' found")


class TestPageBreaks:
    def test_page_break_before_heading(self, converter, simple_template_bytes, basic_mapping_rules):
        rules = {**basic_mapping_rules, "page_break_before": ["heading.2"]}
        md = "# Title\n\n## Chapter\n\nContent"
        docx_bytes, _ = converter.convert(md, simple_template_bytes, rules)
        doc = Document(io.BytesIO(docx_bytes))
        assert len(doc.paragraphs) > 0


class TestMissingStyles:
    def test_missing_style_falls_back_gracefully(self, converter, simple_template_bytes):
        rules = {
            "heading": {"1": "NonExistent Style"},
            "paragraph": "Also Missing",
        }
        docx_bytes, report = converter.convert("# Title\n\nText", simple_template_bytes, rules)
        assert isinstance(docx_bytes, bytes)
        assert len(report["warnings"]) > 0

    def test_empty_mapping_rules(self, converter, simple_template_bytes):
        docx_bytes, report = converter.convert("# Title\n\nText", simple_template_bytes, {})
        assert isinstance(docx_bytes, bytes)


class TestCoverPageTemplate:
    def test_cover_title_updated_from_h1(self, converter, cover_page_template_bytes, cover_page_mapping_rules):
        md = "# My Report\n\nContent here"
        docx_bytes, _ = converter.convert(md, cover_page_template_bytes, cover_page_mapping_rules)
        paras = read_docx_paragraphs(docx_bytes)
        cover_titles = [p for p in paras if p["style"] == "Cover heading"]
        assert any("My Report" in p["text"] for p in cover_titles)

    def test_cover_subtitle_cleared_auto_mode(self, converter, cover_page_template_bytes, cover_page_mapping_rules):
        md = "# Title\n\nBody"
        docx_bytes, _ = converter.convert(md, cover_page_template_bytes, cover_page_mapping_rules)
        paras = read_docx_paragraphs(docx_bytes)
        subtitle_paras = [p for p in paras if p["style"] == "Cover subtitle"]
        for p in subtitle_paras:
            assert p["text"].strip() == ""

    def test_body_content_after_cover(self, converter, cover_page_template_bytes, cover_page_mapping_rules):
        md = "# Title\n\nBody paragraph"
        docx_bytes, _ = converter.convert(md, cover_page_template_bytes, cover_page_mapping_rules)
        paras = read_docx_paragraphs(docx_bytes)
        body_texts = [p["text"] for p in paras if p["style"] == "Normal"]
        assert "Body paragraph" in body_texts

    def test_h1_not_duplicated_in_body(self, converter, cover_page_template_bytes, cover_page_mapping_rules):
        """When H1 updates cover title, it shouldn't also appear in the body."""
        md = "# Title\n\n## Section\n\nContent"
        docx_bytes, report = converter.convert(md, cover_page_template_bytes, cover_page_mapping_rules)
        paras = read_docx_paragraphs(docx_bytes)
        title_paras = [p for p in paras if p["text"] == "Title"]
        assert len(title_paras) <= 1


class TestThreeSectionTemplate:
    def test_produces_valid_docx(self, converter, three_section_template_bytes, basic_mapping_rules):
        md = "# Title\n\nBody content\n\n## Section 2\n\nMore content"
        docx_bytes, _ = converter.convert(md, three_section_template_bytes, basic_mapping_rules)
        doc = Document(io.BytesIO(docx_bytes))
        assert len(doc.paragraphs) > 0

    def test_body_content_inserted(self, converter, three_section_template_bytes, basic_mapping_rules):
        md = "# Title\n\nInserted content here"
        docx_bytes, _ = converter.convert(md, three_section_template_bytes, basic_mapping_rules)
        paras = read_docx_paragraphs(docx_bytes)
        texts = [p["text"] for p in paras]
        assert "Inserted content here" in texts


class TestConversionReport:
    def test_report_has_all_stat_keys(self, converter, simple_template_bytes, basic_mapping_rules):
        _, report = converter.convert("# H\n\nP\n\n- L", simple_template_bytes, basic_mapping_rules)
        stats = report["stats"]
        for key in ("headings", "paragraphs", "lists", "tables", "code_blocks", "images"):
            assert key in stats

    def test_report_internal_keys_excluded(self, converter, simple_template_bytes, basic_mapping_rules):
        _, report = converter.convert("# H", simple_template_bytes, basic_mapping_rules)
        stats = report["stats"]
        assert "_first_h1_done" not in stats

    def test_report_warnings_text_format(self, converter, simple_template_bytes):
        rules = {"heading": {"1": "Missing Style"}}
        _, report = converter.convert("# H", simple_template_bytes, rules)
        assert isinstance(report["warnings_text"], list)
        if report["warnings_text"]:
            assert isinstance(report["warnings_text"][0], str)

    def test_elements_processed_is_sum_of_stats(self, converter, simple_template_bytes, basic_mapping_rules):
        md = "# Title\n\nParagraph\n\n- List"
        _, report = converter.convert(md, simple_template_bytes, basic_mapping_rules)
        assert report["elements_processed"] == sum(report["stats"].values())


class TestTableOfContents:
    def test_toc_inserted_when_mapping_enabled(self, converter, simple_template_bytes, basic_mapping_rules):
        rules = {**basic_mapping_rules, "toc": True}
        md = "# Chapter 1\n\nContent\n\n## Section 1.1\n\nMore"
        docx_bytes, _ = converter.convert(md, simple_template_bytes, rules)
        doc = Document(io.BytesIO(docx_bytes))
        body_xml = doc.element.body.xml
        assert "TOC" in body_xml, "No TOC field code found in document body"
        assert "fldChar" in body_xml, "No field character elements found"

    def test_toc_not_inserted_when_disabled(self, converter, simple_template_bytes, basic_mapping_rules):
        md = "# Chapter 1\n\nContent"
        docx_bytes, _ = converter.convert(md, simple_template_bytes, basic_mapping_rules)
        doc = Document(io.BytesIO(docx_bytes))
        body_xml = doc.element.body.xml
        assert "TOC" not in body_xml

    def test_toc_update_fields_set(self, converter, simple_template_bytes, basic_mapping_rules):
        rules = {**basic_mapping_rules, "toc": True}
        md = "# Title\n\n## Section"
        docx_bytes, _ = converter.convert(md, simple_template_bytes, rules)
        doc = Document(io.BytesIO(docx_bytes))
        settings_xml = doc.settings.element.xml
        assert "updateFields" in settings_xml


class TestHyperlinks:
    def test_link_is_clickable(self, converter, simple_template_bytes, basic_mapping_rules):
        md = "Visit [Example](https://example.com) for info"
        docx_bytes, _ = converter.convert(md, simple_template_bytes, basic_mapping_rules)
        doc = Document(io.BytesIO(docx_bytes))
        from docx.oxml.ns import qn
        hyperlinks = []
        for para in doc.paragraphs:
            for elem in para._p:
                if elem.tag == qn("w:hyperlink"):
                    hyperlinks.append(elem)
        assert len(hyperlinks) >= 1, "No w:hyperlink element found"
        r_id = hyperlinks[0].get(qn("r:id"))
        assert r_id is not None, "Hyperlink missing r:id relationship"

    def test_link_text_preserved(self, converter, simple_template_bytes, basic_mapping_rules):
        md = "Click [here](https://example.com)"
        docx_bytes, _ = converter.convert(md, simple_template_bytes, basic_mapping_rules)
        paras = read_docx_paragraphs(docx_bytes)
        texts = [p["text"] for p in paras]
        assert any("here" in t for t in texts)

    def test_multiple_links_in_paragraph(self, converter, simple_template_bytes, basic_mapping_rules):
        md = "See [A](https://a.com) and [B](https://b.com)"
        docx_bytes, _ = converter.convert(md, simple_template_bytes, basic_mapping_rules)
        doc = Document(io.BytesIO(docx_bytes))
        from docx.oxml.ns import qn
        hyperlinks = []
        for para in doc.paragraphs:
            for elem in para._p:
                if elem.tag == qn("w:hyperlink"):
                    hyperlinks.append(elem)
        assert len(hyperlinks) >= 2


class TestImageInsertion:
    def test_inline_image_from_bytes(self, converter, simple_template_bytes, basic_mapping_rules):
        """Test the internal _add_image_to_paragraph method."""
        import base64
        tiny_png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR4nGP4z8BQDwAEgAF/pooBPQAAAABJRU5ErkJggg=="
        )
        doc = Document(io.BytesIO(simple_template_bytes))
        para = doc.add_paragraph()
        converter._add_image_to_paragraph(para, tiny_png, "test", width_inches=2.0)
        from docx.oxml.ns import qn
        drawings = para._p.findall(f".//{qn('w:drawing')}")
        assert len(drawings) >= 1, "No drawing element found"

    def test_image_placeholder_for_failed_download(self, converter, simple_template_bytes, basic_mapping_rules):
        md = "![broken](https://nonexistent.invalid/img.png)"
        docx_bytes, _ = converter.convert(md, simple_template_bytes, basic_mapping_rules)
        paras = read_docx_paragraphs(docx_bytes)
        texts = [p["text"] for p in paras]
        assert any("broken" in t for t in texts)

    def test_base64_data_uri_image(self, converter, simple_template_bytes, basic_mapping_rules):
        """Test base64 data URI images are inserted."""
        # 1x1 red PNG as data URI
        data_uri = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR4nGP4z8BQDwAEgAF/pooBPQAAAABJRU5ErkJggg=="
        md = f"![red pixel]({data_uri})"
        docx_bytes, report = converter.convert(md, simple_template_bytes, basic_mapping_rules)
        doc = Document(io.BytesIO(docx_bytes))
        from docx.oxml.ns import qn
        drawings = doc.element.body.findall(f".//{qn('w:drawing')}")
        assert len(drawings) >= 1, "No drawing element found for base64 image"
        assert report["stats"]["images"] >= 1
