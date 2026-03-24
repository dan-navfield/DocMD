"""Shared test fixtures for all test modules."""
import io
import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


WP_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


@pytest.fixture
def simple_template_bytes():
    """Single-section template with basic styles (Normal, Heading 1-3, Code, Quote)."""
    doc = Document()

    for i in range(1, 4):
        name = f"Heading {i}"
        if name not in [s.name for s in doc.styles]:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    for name in ("Code", "Quote", "List Bullet", "List Number"):
        if name not in [s.name for s in doc.styles]:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    doc.add_paragraph("Template placeholder", "Normal")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def cover_page_template_bytes():
    """Two-section template: cover page with section break, then body."""
    doc = Document()

    for name in ("Cover heading", "Cover subtitle"):
        doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    for i in range(1, 4):
        name = f"Heading {i}"
        if name not in [s.name for s in doc.styles]:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    for name in ("Code", "Quote", "List Bullet", "List Number"):
        if name not in [s.name for s in doc.styles]:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    doc.add_paragraph("Document Title", "Cover heading")
    doc.add_paragraph("Subtitle here", "Cover subtitle")

    doc.add_section()

    doc.add_paragraph("Body placeholder", "Normal")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def three_section_template_bytes():
    """Three-section template: cover + body + final page."""
    doc = Document()

    for name in ("Cover heading", "Cover subtitle"):
        doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    for i in range(1, 4):
        name = f"Heading {i}"
        if name not in [s.name for s in doc.styles]:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    for name in ("Code", "Quote", "List Bullet", "List Number"):
        if name not in [s.name for s in doc.styles]:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    doc.add_paragraph("Title", "Cover heading")
    doc.add_section()

    doc.add_paragraph("Body content", "Normal")
    doc.add_section()

    doc.add_paragraph("Appendix", "Normal")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def basic_mapping_rules():
    """Mapping rules that match the simple_template_bytes styles."""
    return {
        "heading": {"1": "Heading 1", "2": "Heading 2", "3": "Heading 3"},
        "document_title": "",
        "document_subtitle": "",
        "paragraph": "Normal",
        "list_bullet": "List Bullet",
        "list_bullet_2": "List Bullet",
        "list_bullet_3": "List Bullet",
        "list_ordered": "List Number",
        "list_ordered_2": "List Number",
        "list_ordered_3": "List Number",
        "code_block": "Code",
        "blockquote": "Quote",
        "table": {"style": "Table Grid", "header_row": True},
        "page_break_before": [],
        "metadata_mapping": {},
    }


@pytest.fixture
def cover_page_mapping_rules():
    """Mapping rules for cover page template."""
    return {
        "heading": {"1": "Heading 1", "2": "Heading 2", "3": "Heading 3"},
        "document_title": "Cover heading",
        "document_subtitle": "Cover subtitle",
        "paragraph": "Normal",
        "list_bullet": "List Bullet",
        "list_bullet_2": "List Bullet",
        "list_bullet_3": "List Bullet",
        "list_ordered": "List Number",
        "list_ordered_2": "List Number",
        "list_ordered_3": "List Number",
        "code_block": "Code",
        "blockquote": "Quote",
        "table": {"style": "Table Grid", "header_row": True},
        "page_break_before": [],
        "metadata_mapping": {},
    }


def read_docx_paragraphs(docx_bytes: bytes) -> list[dict]:
    """Helper: read a docx and return list of {text, style} dicts."""
    doc = Document(io.BytesIO(docx_bytes))
    return [
        {"text": p.text, "style": p.style.name if p.style else "Normal"}
        for p in doc.paragraphs
    ]


def read_docx_tables(docx_bytes: bytes) -> list[list[list[str]]]:
    """Helper: read tables from docx as list of rows of cells."""
    doc = Document(io.BytesIO(docx_bytes))
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append(rows)
    return tables
